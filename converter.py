import os
import shutil

from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement

import docx
import openpyxl
from docx import Document
from docx.shared import Inches, Pt
from docxcompose.composer import Composer

from utils.PathManager import load_path_manager as lpm

file_prog_spec_template = lpm.template("prog_spec_template.docx")
file_prog_spec_word_template_part1 = lpm.template("prog_spec_word_template_part1.docx")
file_prog_spec_word_template_part2 = lpm.template("prog_spec_word_template_part2.docx")
file_prog_spec_word_template_part3 = lpm.template("prog_spec_word_template_part3.docx")
file_prog_spec_output_template = lpm.template("prog_spec_output_template.docx")
file_prog_spec_output_raw_template = lpm.template("prog_spec_output_raw_template.docx")
file_output = lpm.output("output.docx")
file_output_raw = lpm.output("output_raw.docx")
file_output_api_unit = lpm.transit.converter("output_api_unit.docx")
file_output_module_unit = lpm.transit.converter("output_module_unit.docx")

folder_output = str(lpm.output)
folder_transit_creator = str(lpm.transit.creator)
folder_input_logical_sequences = str(lpm.input("Logical Sequences"))


def read_excel_column_headers(excel_path, sheet_name):
    wb = openpyxl.load_workbook(excel_path)
    ws = wb[sheet_name]
    column_headers = [cell.value for cell in ws[1]]
    wb.close()
    return column_headers


# Deprecated
def read_excel_sheet(excel_path, sheet_name):
    wb = openpyxl.load_workbook(excel_path)
    ws = wb[sheet_name]
    data = []
    for row in ws.iter_rows(min_row=2, values_only=True):
        data.append(row)
    wb.close()
    return data


def fill_word_template(file_api_unit, data_dict):
    doc = Document(file_api_unit)
    for p in doc.paragraphs:
        for key, value in data_dict.items():
            p.text = p.text.replace("{{" + key + "}}", value)
    return doc


def read_excel_table(excel_path, sheet_name):
    wb = openpyxl.load_workbook(excel_path)
    ws = wb[sheet_name]
    table_data = []
    for row in ws.iter_rows(values_only=True):
        if any(cell is not None for cell in row):
            table_data.append(row)
    wb.close()
    return table_data


def set_table_column_width(table, num_columns):
    column_percentages = []
    if num_columns == 2:
        column_percentages = [30, 70]
    elif num_columns == 3:
        column_percentages = [20, 60, 20]
    elif num_columns == 4:
        column_percentages = [20, 20, 40, 20]
    elif num_columns == 5:
        column_percentages = [15, 15, 15, 25, 30]
    else:
        pass

    page_width_in_points = 8.5 * 72
    column_widths = [page_width_in_points * percentage / 100 for percentage in column_percentages]

    for idx, width in enumerate(column_widths):
        column_width_in_points = round(width * 10)
        table.columns[idx].width = column_width_in_points


def fill_table_template(temp_path, sheet_name, table_data):
    doc = Document(temp_path)

    # fill table data
    if table_data:
        for table in doc.tables:
            # print(table.cell(0, 0).text)
            if f"{{{sheet_name}}}" in table.cell(0, 0).text:

                for row_index, row_data in enumerate(table_data):

                    if row_index == 0:
                        continue  # skip set header value

                    row = table.add_row().cells

                    for col_index, cell_data in enumerate(row_data):
                        cell = row[col_index]
                        cell.text = str(cell_data)
                        paragraphs = cell.paragraphs
                        for paragraph in paragraphs:
                            for run in paragraph.runs:
                                font = run.font
                                font.size = Pt(6)

                # remove the first line used as a marker
                first_row = table.rows[0]
                tbl = table._tbl
                tbl.remove(first_row._tr)

                set_table_column_width(table, len(table.columns))

    return doc


def clear_empty_tables(doc):
    tables_to_remove = []
    for table in doc.tables:
        if len(table.rows) == 1 and len(table.columns) != 1:
            # rows == 1 and columns != 1, is header row, indicate the table is empty, so remove it
            tables_to_remove.append(table)
        if len(table.rows) == 2 and table.rows[0].cells[0].text.strip() == "Sample Output":
            # special handling remove sample output mark row
            first_row = table.rows[0]
            tbl = table._tbl
            tbl.remove(first_row._tr)

    for table in tables_to_remove:
        first_row = table.rows[0]
        tbl = table._tbl
        tbl.remove(first_row._tr)


def get_logical_sequences_document(api_id):
    file_list = os.listdir(folder_input_logical_sequences)

    matching_files = [filename for filename in file_list if api_id in filename]

    if not matching_files:
        print(f"   -> no matching Logical Sequences file found")
        return Document()
    else:
        file_logical_sequences = Document(os.path.join(folder_input_logical_sequences, matching_files[0]))
        return file_logical_sequences


def generate_single_api_unit(excel_file_path):
    wb = openpyxl.load_workbook(excel_file_path)
    for sheet_name in wb.sheetnames:

        if sheet_name == "MASTER":

            column_headers = read_excel_column_headers(excel_file_path, sheet_name)
            data_dict = {}

            for header in column_headers:
                data_dict[header] = ""

            ws = wb[sheet_name]
            for col_index, header in enumerate(column_headers, start=1):
                data_dict[header] = str(ws.cell(row=2, column=col_index).value)

            # 1. get logical sequences section
            file_logical_sequences_part = get_logical_sequences_document(data_dict["API ID"])

            # 2. compose template file and append the logical sequences section
            composer = Composer(Document())
            composer.append(Document(file_prog_spec_word_template_part1))
            composer.append(Document(file_prog_spec_word_template_part2))
            composer.append(file_logical_sequences_part)
            composer.append(Document(file_prog_spec_word_template_part3))
            composer.save(file_output_api_unit)

            # 3. fill constant data
            filled_doc = fill_word_template(file_output_api_unit, data_dict)
            filled_doc.save(file_output_api_unit)

            # 4. fill sample output (table data)
            # parse sample output str data to table data
            table_data = [
                tuple(["Sample Output"]),
                tuple([data_dict["Sample Output"]])
            ]

            filled_doc = fill_table_template(file_output_api_unit, "Sample Output", table_data)
            filled_doc.save(file_output_api_unit)

        else:

            # 5. fill others table data
            table_data = read_excel_table(excel_file_path, sheet_name)
            filled_doc = fill_table_template(file_output_api_unit, sheet_name, table_data)

            filled_doc.save(file_output_api_unit)

    wb.close()

    filled_doc = Document(file_output_api_unit)

    # 6. clear template default N/A mark
    clear_template_mark(filled_doc)
    # 7. clear empty table
    clear_empty_tables(filled_doc)
    # 8. bold text style
    bold_text_style(filled_doc)

    filled_doc.save(file_output_api_unit)
    return filled_doc


# Deprecated
def count_elements_checker(doc):
    # Note: doc.element.body count = paragraphs element count + table element count
    paragraph_count = len(doc.paragraphs)
    table_count = len(doc.tables)
    table_count_index = table_count
    current_count = 0
    other_element_count = 0
    table_row_indices = {}
    na_paragraph_indices = []

    for i, element in enumerate(doc.element.body):
        current_count += 1
        if isinstance(element, docx.oxml.text.paragraph.CT_P):
            # convert the element to a paragraph and check if it contains "N/A"
            paragraph = docx.text.paragraph.Paragraph(element, doc)
            if "N/A" in paragraph.text:
                na_paragraph_indices.append(current_count)
        elif isinstance(element, docx.oxml.table.CT_Tbl):
            table_row_indices[table_count_index - 1] = current_count
            table_count_index -= 1
        else:
            other_element_count += 1
            print(f"Index of other element: {current_count}")

    print(f"Number of paragraphs: {paragraph_count}")
    print(f"Number of tables: {table_count}")
    print(f"Number of other elements: {other_element_count}")
    print(f"Table row indices: {table_row_indices}")
    print(f"Paragraph indices with 'N/A': {na_paragraph_indices}")


def init_output_folder():
    # clear all files in folder_output
    for filename in os.listdir(folder_output):
        file_path = os.path.join(folder_output, filename)
        try:
            if os.path.isfile(file_path):
                os.unlink(file_path)
        except Exception as e:
            print(f"Error deleting {file_path}: {e}")

    # copy output_template.docx to folder_output, and rename to output.docx
    shutil.copy(file_prog_spec_output_template, file_output)
    shutil.copy(file_prog_spec_output_raw_template, file_output_raw)


def clear_template_mark(doc):
    # remove the N/A mark which exists in a row on the table
    table_count = len(doc.tables)
    table_count_index = table_count

    current_count = 0

    full_table_indices = {}
    na_paragraph_indices = []
    for i, element in enumerate(doc.element.body):

        if isinstance(element, docx.oxml.text.paragraph.CT_P):
            # convert the element to a paragraph and check if it contains "N/A"
            paragraph = docx.text.paragraph.Paragraph(element, doc)
            if "N/A" in paragraph.text:
                na_paragraph_indices.append(current_count)
        elif isinstance(element, docx.oxml.table.CT_Tbl):
            # convert the element to a table and get the row count
            table = docx.table.Table(element, doc)
            row_count = len(table.rows)
            # print(row_count)

            if row_count > 1:
                # rows == 1, is header row, indicate the table is empty, after will remove it
                # row > 1 , is not empty table , so remove N/A mark
                full_table_indices[table_count_index - 1] = current_count
                table_count_index -= 1
        else:
            continue

        current_count += 1

    # print(full_table_indices)
    # print(na_paragraph_indices)

    for value in reversed(full_table_indices.values()):
        # using reverse remove to avoid index changes
        check_index = value - 1  # mark on a row above the table
        if check_index in na_paragraph_indices:
            # mapping requires removing N/A marks
            paragraph_to_remove = doc.element.body[check_index]

            paragraph = docx.text.paragraph.Paragraph(paragraph_to_remove, doc)
            paragraph_element = paragraph._element
            parent = paragraph_element.getparent()
            parent.remove(paragraph_element)


def bold_text_style(doc):
    target_strings = [
        'API Details',
        'Path Parameter(s)',
        'Query Parameter(s)',
        'Header Parameter(s)',
        'Request Body',
        'Validation',
        'Access Control',
        'Logical Sequences',
        'Data Access Layer'
    ]

    for paragraph in doc.paragraphs:
        for target_string in target_strings:
            if paragraph.text.strip() == target_string:
                for run in paragraph.runs:
                    run.font.bold = True


def add_page_break_marker(file_path):
    doc = Document(file_path)

    page_break_count = 0

    for paragraph in doc.paragraphs:
        if "{{page_break}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{page_break}}", "")
            run = paragraph.add_run()
            run.add_break(docx.enum.text.WD_BREAK.PAGE)
            page_break_count += 1

    doc.save(file_path)
    # print(f"total_page_break_count {total_page_break_count}")


def sort_modules(list):
    sorted_list = sorted(list)
    sorted_list = sorted(sorted_list, key=lambda x: not x.startswith("FS"))

    return sorted_list


def start():
    print(">>>>>>>>>>> CONVERTER PROCESS START >>>>>>>>>>>")
    init_output_folder()
    output_doc_raw = Document()

    for title_folder_name in sort_modules(os.listdir(folder_transit_creator)):
        title_folder_path = os.path.join(folder_transit_creator, title_folder_name)

        if os.path.isdir(title_folder_path):

            module_unit_doc = Document()

            # restore file names containing special characters
            origin_title_folder_name = title_folder_name.replace("_", "/")
            module_unit_doc.add_heading(origin_title_folder_name, level=1)
            composer_module_unit = Composer(module_unit_doc)

            print(origin_title_folder_name)

            # compose api unit to module unit
            for excel_file_name in sorted(os.listdir(title_folder_path)):
                if excel_file_name.endswith(".xlsx"):
                    print("-   " + excel_file_name)
                    excel_file_path = os.path.join(title_folder_path, excel_file_name)

                    generate_single_api_unit(excel_file_path)

                    api_unit_doc = Document(file_output_api_unit)
                    composer_module_unit.append(api_unit_doc)

            composer_module_unit.save(file_output_module_unit)

            # output extend template style
            # NOTE sort layout issue, can manually fix , see refer/Chore/fix-out-put-sort-issue-*.png
            output_doc = Document(file_output)
            module_unit_doc = Document(file_output_module_unit)

            composer = Composer(output_doc)
            composer.append(module_unit_doc)  # append to the docs end by default
            composer.save(file_output)

            # output raw style
            # NOTE sort layout great,
            # but need manually copy and paste to the baseline word(prog_spec_output_template.docx)
            composer = Composer(output_doc_raw)
            composer.append(module_unit_doc)
            composer.save(file_output_raw)

    # add page break
    add_page_break_marker(file_output)
    add_page_break_marker(file_output_raw)

    print(">>>>>>>>>>> CONVERTER PROCESS END >>>>>>>>>>>>>")


if __name__ == '__main__':
    start()
